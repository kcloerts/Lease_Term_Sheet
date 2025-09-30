import streamlit as st
import PyPDF2
from docx import Document
import io
import os
import google.generativeai as genai
from html.parser import HTMLParser

# Set page configuration
st.set_page_config(
    page_title="Lease Term Sheet Generator",
    page_icon="📄",
    layout="wide"
)

copilot/fix-ac37161a-44b2-4617-bc87-5e767d6d98ed
class HTMLTextExtractor(HTMLParser):
    """Extract text content from HTML"""
    def __init__(self):
        super().__init__()
        self.text = []
        self.skip_tags = set()
        
    def handle_starttag(self, tag, attrs):
        # Skip content inside style and script tags
        if tag in ('style', 'script'):
            self.skip_tags.add(tag)
            
    def handle_endtag(self, tag):
        # Re-enable content extraction when closing style/script tags
        self.skip_tags.discard(tag)
        
    def handle_data(self, data):
        # Only add data if we're not inside a skip tag
        if not self.skip_tags and data.strip():
            self.text.append(data.strip())
        
    def get_text(self):
        return '\n'.join(self.text)

def read_html(file):
    """Extract text from HTML file"""
    if isinstance(file, str):
        # File path
        with open(file, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
    else:
        # File-like object
        html_content = file.read()
        if isinstance(html_content, bytes):
            html_content = html_content.decode('utf-8', errors='ignore')
    
    parser = HTMLTextExtractor()
    parser.feed(html_content)
    return parser.get_text()

main

def read_pdf(file):
    """Extract text from PDF file"""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx(file):
    """Extract text from DOCX file"""
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_document(file):
    """Read document based on file type"""
    if file.name.endswith('.pdf'):
        return read_pdf(file)
    elif file.name.endswith('.docx'):
        return read_docx(file)
    elif file.name.endswith('.htm') or file.name.endswith('.html'):
        return read_html(file)
    else:
        return file.read().decode('utf-8')

def create_docx_from_text(text):
    """Create a DOCX document from text"""
    doc = Document()
    
    # Split text into lines and add to document
    lines = text.split('\n')
    for line in lines:
        doc.add_paragraph(line)
    
    # Save to BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file.getvalue()

def list_available_models(api_key):
    """List available Gemini models"""
    try:
        genai.configure(api_key=api_key)
        models = []
        for model in genai.list_models():
            if 'generateContent' in model.supported_generation_methods:
                models.append(model.name)
        return models
    except Exception as e:
        return []

# Load default template from file
def load_default_template():
    """Load the default template from Term Sheet Template_app.htm"""
    template_path = os.path.join(os.path.dirname(__file__), "Term Sheet Template_app.htm")
    try:
        if os.path.exists(template_path):
            return read_html(template_path)
        else:
            # Fallback to a basic template if file not found
            return """COMMERCIAL LEASE TERM SHEET

Property Address: [Address]
Tenant Name: [Tenant Name]
Landlord Name: [Landlord Name]

LEASE TERMS:

1. PREMISES
   - Suite/Unit Number: [Suite]
   - Rentable Square Feet: [SF]
   - Use: [Permitted Use]

2. LEASE TERM
   - Commencement Date: [Date]
   - Expiration Date: [Date]
   - Term Length: [Years/Months]
   - Option to Extend: [Yes/No, Terms]

3. BASE RENT
   - Initial Annual Base Rent: [Amount]
   - Monthly Base Rent: [Amount]
   - Rent Escalations: [Schedule]

4. ADDITIONAL RENT
   - Operating Expenses: [Details]
   - Property Taxes: [Details]
   - Utilities: [Responsibility]
   - CAM Charges: [Details]

5. SECURITY DEPOSIT
   - Amount: [Amount]
   - Terms: [Details]

6. TENANT IMPROVEMENTS
   - Tenant Improvement Allowance: [Amount]
   - Construction Period: [Timeline]

7. PARKING
   - Number of Spaces: [Number]
   - Type: [Reserved/Unreserved]
   - Cost: [Amount if any]

8. SPECIAL PROVISIONS
   - [Any special terms or conditions]

9. BROKER INFORMATION
   - Landlord's Broker: [Name]
   - Tenant's Broker: [Name]
"""
    except Exception as e:
        # Return fallback template if there's an error
        return """COMMERCIAL LEASE TERM SHEET

Property Address: [Address]
Tenant Name: [Tenant Name]
Landlord Name: [Landlord Name]

LEASE TERMS:

1. PREMISES
2. LEASE TERM
3. BASE RENT
4. ADDITIONAL RENT
5. SECURITY DEPOSIT
6. TENANT IMPROVEMENTS
7. PARKING
8. SPECIAL PROVISIONS
9. BROKER INFORMATION
"""

DEFAULT_TEMPLATE = load_default_template()

def generate_term_sheet(template_text, lease_text, api_key):
    """Generate term sheet using Gemini API"""
    
    prompt = f"""You are a commercial real estate expert. You have been provided with:
1. A lease term sheet template
2. A full commercial lease document

Your task is to analyze the commercial lease and extract all relevant information to create a completed lease term sheet that matches the template format exactly.

LEASE TERM SHEET TEMPLATE:
{template_text}

COMMERCIAL LEASE:
{lease_text}

Please generate a completed lease term sheet that:
1. Follows the exact structure and format of the template
2. Extracts all relevant information from the commercial lease
3. Fills in all sections of the template with appropriate data from the lease
4. Maintains professional formatting
5. Uses clear, concise language
6. If information is not found in the lease, indicate "Not specified in lease"

Generate the completed lease term sheet now:"""

    try:
        genai.configure(api_key=api_key)
        # Use gemini-2.5-pro for advanced lease analysis capabilities
        model = genai.GenerativeModel('gemini-2.5-pro')
        
        full_prompt = f"""You are an expert commercial real estate attorney specializing in lease analysis and term sheet creation.

{prompt}"""
        
        response = model.generate_content(
            full_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=4000,
            )
        )
        return response.text
        
    except Exception as e:
        error_msg = str(e)
        # If model not found, try to list available models
        if "not found" in error_msg.lower() or "not supported" in error_msg.lower():
            available_models = list_available_models(api_key)
            if available_models:
                models_str = "\n".join([f"  - {m}" for m in available_models])
                return f"Error: The specified model is not available.\n\nAvailable models that support content generation:\n{models_str}\n\nOriginal error: {error_msg}"
            else:
                return f"Error generating term sheet: {error_msg}\n\nTip: Common model names include 'gemini-2.5-pro', 'gemini-1.5-flash', 'gemini-1.5-pro', or 'gemini-pro'"
        return f"Error generating term sheet: {error_msg}"

# Main app
def main():
    st.title("📄 Lease Term Sheet Generator")
    st.markdown("""
    This application helps you generate a lease term sheet by analyzing a commercial lease document.
    A default template is provided, or you can upload your own custom template.
    """)
    
    # API Key configuration
    st.sidebar.header("Configuration")
    
    # Try to get default API key from secrets
    default_api_key = None
    try:
        default_api_key = st.secrets.get('gemini', {}).get('api_key', None)
    except Exception:
        pass
    
    # API Key input - optional if default is available
    if default_api_key:
        st.sidebar.success("✅ Using default Gemini API key")
        use_custom_key = st.sidebar.checkbox("Use custom API key", value=False)
        if use_custom_key:
            api_key = st.sidebar.text_input(
                "Google Gemini API Key", 
                type="password",
                help="Enter your Google Gemini API key to override the default"
            )
            if not api_key:
                api_key = default_api_key
        else:
            api_key = default_api_key
    else:
        st.sidebar.info("💡 No default API key configured")
        api_key = st.sidebar.text_input(
            "Google Gemini API Key",
            type="password",
            help="Enter your Google Gemini API key to enable AI-powered analysis"
        )
    
    if not api_key:
        st.warning("⚠️ Please enter your Google Gemini API key in the sidebar to use this application.")
        st.info("""
        To use this application:
        1. Get an API key from [Google AI Studio](https://makersuite.google.com/app/apikey)
        2. Enter it in the sidebar
        3. Upload your documents
        """)
        return
    
    # Create two columns for file uploads
    col1, col2 = st.columns(2)
    
    with col1:
 copilot/fix-ac37161a-44b2-4617-bc87-5e767d6d98ed
        st.subheader("1️⃣ Template")
        use_custom_template = st.checkbox("Use custom template", value=False, 
                                          help="Check this to upload your own template instead of using the default")
        
        if use_custom_template:
            template_file = st.file_uploader(
                "Upload template (PDF, DOCX, TXT, or HTM)",
                type=['pdf', 'docx', 'txt', 'htm', 'html'],

     
 main
                key="template",
                help="Upload your lease term sheet template that will be used as the format"
            )
            
            if template_file:
                st.success(f"✅ Template uploaded: {template_file.name}")
        else:
            st.info("✅ Using default template")
            with st.expander("📄 View Default Template"):
copilot/fix-ac37161a-44b2-4617-bc87-5e767d6d98ed
                st.text_area("Default Template Content", DEFAULT_TEMPLATE[:1000] + "..." if len(DEFAULT_TEMPLATE) > 1000 else DEFAULT_TEMPLATE, height=300, disabled=True)


main
    
    with col2:
        st.subheader("2️⃣ Upload Commercial Lease")
        lease_file = st.file_uploader(
            "Upload lease (PDF, DOCX, or TXT)",
            type=['pdf', 'docx', 'txt'],
            key="lease",
            help="Upload the commercial lease document to analyze"
        )
        
        if lease_file:
            st.success(f"✅ Lease uploaded: {lease_file.name}")
    
    # Process documents when lease is uploaded
    if lease_file:
        st.markdown("---")
        
        if st.button("🚀 Generate Term Sheet", type="primary"):
            with st.spinner("Reading documents..."):
                try:
                    # Get template text (use custom or default)
                    if use_custom_template and template_file:
                        template_text = read_document(template_file)
                    else:
                        template_text = DEFAULT_TEMPLATE
                    
                    # Read lease document
                    lease_text = read_document(lease_file)
                    
                    st.success("✅ Documents read successfully!")
                    
                    # Show preview in expanders
                    with st.expander("📄 View Template Preview"):
                        st.text_area("Template Content", template_text[:2000] + "..." if len(template_text) > 2000 else template_text, height=200, disabled=True)
                    
                    with st.expander("📄 View Lease Preview"):
                        st.text_area("Lease Content", lease_text[:2000] + "..." if len(lease_text) > 2000 else lease_text, height=200, disabled=True)
                    
                except Exception as e:
                    st.error(f"❌ Error reading documents: {str(e)}")
                    return
            
            with st.spinner("🤖 Analyzing lease and generating term sheet... This may take 30-60 seconds."):
                try:
                    # Generate term sheet
                    term_sheet = generate_term_sheet(template_text, lease_text, api_key)
                    
                    st.success("✅ Term sheet generated successfully!")
                    
                    # Display result
                    st.markdown("---")
                    st.subheader("📋 Generated Lease Term Sheet")
                    st.markdown(term_sheet)
                    
                    # Download button
                    docx_data = create_docx_from_text(term_sheet)
                    st.download_button(
                        label="⬇️ Download Term Sheet",
                        data=docx_data,
                        file_name="lease_term_sheet.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"❌ Error generating term sheet: {str(e)}")
    else:
        st.info("👆 Please upload a lease document to begin.")

if __name__ == "__main__":
    main()
